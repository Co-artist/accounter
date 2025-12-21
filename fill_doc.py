import sys
from docx import Document
import os

docx_path = r"C:\Users\zzh66\Desktop\“创想之星”系列——软件著作大赛提交模板_20251201162655.docx"

if not os.path.exists(docx_path):
    print(f"File not found: {docx_path}")
    # Try removing quotes from filename if they were part of the path string but not the filename?
    # But LS showed the quotes are part of the filename.
    sys.exit(1)

try:
    doc = Document(docx_path)
except Exception as e:
    print(f"Error opening file: {e}")
    sys.exit(1)

def insert_text_after_header(header_text, text_to_insert):
    found = False
    for paragraph in doc.paragraphs:
        # Check if header is in paragraph text. 
        # The document might use different whitespace or numbering styles.
        # We strip to be safer.
        if header_text in paragraph.text:
            new_p = doc.add_paragraph()
            new_p.add_run(text_to_insert)
            # Move the new paragraph to be after the found paragraph
            paragraph._element.addnext(new_p._element)
            print(f"Inserted text after header containing: {header_text}")
            found = True
            break
    
    if not found:
        print(f"Warning: Header '{header_text}' not found in document.")

# Content
intro_text = """项目名称：猫头鹰记账 (Owl Accounting)

核心理念：
本项目是一款基于 Vue 3 + Capacitor 构建的跨平台个人记账应用。它打破了传统记账软件枯燥、严肃的工具属性，以“情感化设计”为核心，引入了一只可交互的“猫头鹰”吉祥物作为视觉中心，旨在通过温馨、趣味的交互体验，帮助用户养成坚持记账的好习惯。

主要功能与实用性：
1. 全方位账目管理：支持极速记录收入/支出，自定义分类图标与颜色，满足日常收支管理需求。
2. 多维度数据分析：内置 ECharts 强力驱动的统计报表，支持按周、月、年查看收支趋势、分类占比饼图，帮助用户清晰洞察财务状况。
3. 智能预算控制：支持设定月度预算，实时监控剩余额度，防止冲动消费。
4. 安全便捷的云同步：基于 Node.js + MongoDB 构建后端服务，通过 JWT 安全认证实现多端数据实时同步，换机无忧。

创新性与亮点：
1. 趣味交互体验：登录/注册页面的猫头鹰 Logo 具有仿生交互能力——当用户输入密码时，猫头鹰会害羞地用翅膀捂住眼睛，保护隐私同时也增加了趣味性；页面背景伴随动态气泡与光效，打破沉闷。
2. 丝滑的移动端手感：深度优化的单页应用 (SPA) 架构，实现了原生级别的左右滑动切换页面体验，配合精心调试的转场动画，彻底消除了页面跳转的割裂感与白屏现象。
3. 现代化 UI 设计：采用 Glassmorphism（毛玻璃）风格设计语言，配合温暖的黄/棕色主色调，界面清新雅致，视觉层级分明。"""

env_text = """本项目采用前后端分离架构，并支持打包为 Android 原生应用。

开发与运行环境：
操作系统：Windows 10/11, macOS, 或 Linux
移动端运行环境：Android 7.0 及以上版本
Node.js 版本：v16.0.0 及以上 (推荐 v18+)

核心技术栈与库版本：
1. 前端框架：Vue.js ^3.5.13 (Composition API setup 语法糖)
2. 构建工具：Vite ^6.0.5
3. 跨平台容器：Capacitor ^8.0.0 (Android 平台支持)
4. 路由管理：Vue Router ^4.6.3
5. 图表可视化：ECharts ^6.0.0
6. 网络请求：Axios ^1.13.2
7. 动画引擎：GSAP ^3.14.2 (用于复杂动效)
8. 后端服务：Express ^5.2.1 + Mongoose ^9.0.1 (MongoDB 驱动) + JSONWebToken ^9.0.3"""

screenshots_text = "(此处请手动插入截图：登录/注册页、首页、记账页、统计页、个人中心)"

summary_text = """技术成长的跨越：
在开发“猫头鹰记账”的过程中，我深入实践了 Vue 3 的 Composition API 和现代前端工程化流程。最大的挑战在于如何在一个 Web 技术栈的应用中实现媲美原生 App 的流畅度。通过引入 Capacitor，我成功打通了 Web 与 Native 的桥梁。在解决“页面切换白屏”和“转场动画卡顿”等典型混合开发难题时，我深入研究了 Vue Router 的生命周期和 CSS3 硬件加速机制，最终通过预加载策略和绝对定位布局方案完美解决了这些问题。

产品思维的觉醒：
不仅仅是写代码，更是在做产品。我意识到一个好的应用不仅要功能无由于，更要有温度。项目中“猫头鹰捂眼睛”的微交互设计，虽然技术实现不复杂，但却能给用户带来极大的惊喜感，这让我深刻体会到细节体验对于用户留存的重要性。

全栈开发的初探：
从设计数据库模型 (MongoDB Schema) 到编写 RESTful API，再到前端对接与鉴权逻辑的实现，这次完整的全栈开发经历极大地拓宽了我的技术视野，让我能够从更宏观的角度去思考系统的稳定性与安全性。"""

# Matching partial strings to avoid unicode issues with specific numbering dots
insert_text_after_header("程序简介", intro_text)
insert_text_after_header("程序运行环境", env_text)
insert_text_after_header("程序运行截图", screenshots_text)
insert_text_after_header("总结和心得", summary_text)

output_path = docx_path.replace(".docx", "_已填写.docx")
doc.save(output_path)
print(f"Success! Document saved to: {output_path}")
